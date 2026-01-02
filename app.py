from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import fitz
from docx import Document
from groq import Groq
import os
import uuid
import re
from dotenv import load_dotenv
from deep_translator import GoogleTranslator

# ---------------- LOAD ENV ----------------
load_dotenv()

# ---------------- APP SETUP ----------------
app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# ---------------- TRANSLATION ----------------
def ensure_english(text):
    try:
        return GoogleTranslator(source="auto", target="en").translate(text)
    except:
        return text

# ---------------- PDF TEXT ----------------
def extract_full_text(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

# ---------------- SECTION SPLITTER ----------------
def extract_sections(text):
    sections = {"FULL_CONTENT": ""}
    current = "FULL_CONTENT"

    for line in text.split("\n"):
        if re.search(r"(UNIT\s*\d+|CHAPTER\s*\d+)", line, re.I):
            current = line.strip()
            sections[current] = ""
        else:
            sections[current] += line + " "

    return sections

def find_relevant_section(instruction, sections):
    instruction = instruction.lower()
    for title, content in sections.items():
        if title.lower() in instruction:
            return content
    return sections["FULL_CONTENT"]

# ---------------- CREATE DOCX ----------------
def create_doc(content):
    filename = f"questions_{uuid.uuid4().hex}.docx"
    path = os.path.join(OUTPUT_FOLDER, filename)

    doc = Document()
    doc.add_heading("Generated Question Paper", level=1)

    for line in content.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    doc.save(path)
    return filename

# ---------------- ROUTES ----------------
@app.route("/")
def home():
    return render_template("index.html")

# 🔹 TRANSLATION API (FIXED)
@app.route("/translate", methods=["POST"])
def translate():
    data = request.json
    text = data.get("text", "")
    translated = ensure_english(text)
    return jsonify({"translated": translated})

@app.route("/generate", methods=["POST"])
def generate():
    try:
        pdf = request.files.get("pdf")
        instruction = request.form.get("instruction")
        lang = request.form.get("lang", "en")

        if not pdf or not instruction:
            return jsonify({"error": "Missing PDF or instruction"}), 400

        if lang == "ta":
            instruction = ensure_english(instruction)

        pdf_path = os.path.join(UPLOAD_FOLDER, "reference.pdf")
        pdf.save(pdf_path)

        full_text = extract_full_text(pdf_path)
        sections = extract_sections(full_text)
        content = find_relevant_section(instruction, sections)

        prompt = f"""
Generate exam questions in ENGLISH ONLY.
No explanations.

CONTENT:
{content}

TASK:
{instruction}
"""

        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )

        filename = create_doc(response.choices[0].message.content)

        return jsonify({
            "message": "Questions generated successfully!",
            "filename": filename
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/download/<filename>")
def download(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path):
        return "File not found", 404
    return send_file(path, as_attachment=True)

# ---------------- RUN ----------------
if __name__ == "__main__":
    app.run(debug=True)
@app.route("/")
def home():
    return render_template("index.html")
